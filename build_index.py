"""
build_index.py -- QP Insights Commons
Indexes all .txt and .docx files from the REPO ROOT.
Paths stored as raw filenames (no URL-encoding) so Worker can encode correctly.
"""

import json, os, re

OUTPUT_FILE = "index.json"

STOPWORDS = {
    "the","a","an","and","or","but","in","on","at","to","for","of","with",
    "is","are","was","were","be","been","being","have","has","had","do","does",
    "did","will","would","could","should","may","might","shall","can","this",
    "that","these","those","it","its","they","their","we","our","you","your",
    "how","what","when","where","why","which","who","from","by","as","into",
    "per","through","during","before","after","above","below","up","down","out"
}

COOKBOOK_META = {
    "ai_router_technical":                 ("AI Router Technical Setup Guide",
                                            "ai router webhook json custom variable survey prompt setup technical cookbook"),
    "communities":                          ("Purpose Mindset and Communities Cookbook",
                                            "communities purpose mindset personal statement fortune teller ai router"),
    "conversational_survey":               ("Conversational Survey Cookbook",
                                            "conversational survey sous chef adaptive ai router nps csat ces"),
    "conversational_survey_ready_to_sell": ("Conversational Survey Ready to Sell",
                                            "conversational survey sous chef adaptive ai router pricing rts sell benefits"),
    "salesforce_integration":              ("Salesforce Integration Cookbook",
                                            "salesforce smoke alarm triggering survey crm flow named credential callout hpe hewlett packard"),
    "salesforce_ready_to_sell":            ("Salesforce Triggering Ready to Sell",
                                            "salesforce smoke alarm trigger survey crm pricing rts sell benefits"),
    "sentiment_analysis":                  ("Sentiment Analysis Cookbook",
                                            "sentiment analysis prep cook open-ended text ai router sartorius themes summary dashboard"),
    "sentiment_analysis_ready_to_sell":    ("Sentiment Analysis Ready to Sell",
                                            "sentiment analysis prep cook open-ended text pricing rts sell benefits"),
    "tv_guide_customisation":              ("Dynamic TV Guide Searchable List Cookbook",
                                            "tv guide the menu dynamic list middleware searchable catalogue proxy trp research ireland"),
    "tv_guide_ready_to_sell":              ("Dynamic TV Guide Ready to Sell",
                                            "tv guide the menu dynamic list searchable catalogue pricing rts sell benefits"),
    "solution_intro":                      ("Solutions Cookbook Overview Quick Reference",
                                            "solutions overview pricing proof point sous chef prep cook fortune teller smoke alarm the menu"),
    "custom_canvas_dashboards":            ("Custom Canvas Dashboards Guide",
                                            "canvas dashboard widget custom visualization cx insights"),
    "custom_dashboards_ready_to_sell":     ("Custom Dashboards Ready to Sell",
                                            "canvas dashboard widget custom visualization pricing rts sell benefits"),
    "intercept_mobile_sdk___technical_doc":("Intercept Mobile SDK Technical Doc",
                                            "intercept mobile sdk ios android technical setup"),
    "questionpro___middleware_document":   ("QuestionPro Middleware Document",
                                            "middleware proxy api integration custom"),
    "questionpro____intercept____setup_guide":("Intercept Setup Guide",
                                            "intercept setup guide cx survey popup widget trigger data layer rule targeting datalayer"),
}

SKIP_FILES = {"build_index.py","index.json","index.html","worker.js",
              "README.md","help-sitemap.json","build_index_action.yml","KNOWLEDGE_README.md"}

def normalise_key(fn):
    return re.sub(r'[^a-z0-9_]', '_', fn.lower().replace('.txt','').replace('.docx',''))

def extract_keywords(text, n=30):
    text = re.sub(r'[^a-z0-9\s]',' ', text.lower())
    freq = {}
    for w in text.split():
        if len(w)>=2 and w not in STOPWORDS: freq[w]=freq.get(w,0)+1
    return ' '.join(w for w,_ in sorted(freq.items(),key=lambda x:-x[1])[:n])

def read_docx(filename):
    """Extract text from a .docx file."""
    try:
        from docx import Document
        doc = Document(filename)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # Also grab table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return '\n'.join(parts)
    except ImportError:
        # python-docx not available, return empty
        print(f"  WARNING: python-docx not installed, skipping {filename}")
        return ''
    except Exception as e:
        print(f"  WARNING: could not read {filename}: {e}")
        return ''

def index_plaintext(filename):
    with open(filename,'r',encoding='utf-8',errors='replace') as f: content=f.read()
    key = normalise_key(filename)
    meta = COOKBOOK_META.get(key)
    title, extra_kw = meta if meta else (filename.replace('.txt','').replace('.docx',''), '')
    kw = extract_keywords(content) + (' '+extra_kw if extra_kw else '')
    clean = re.sub(r'[+|=\-]{2,}',' ',content)
    clean = re.sub(r'\s+',' ',clean).strip()
    fn_l = filename.lower()
    cat = ('sales' if 'ready to sell' in fn_l or 'ready_to_sell' in fn_l
           else 'technical' if any(x in fn_l for x in ['technical','integration','customis'])
           else 'correction' if fn_l.startswith('correction_')
           else 'general')
    return [{"id":key,"title":title,"category":cat,"product":"CX Solutions",
             "kw":kw,"summary":clean[:300],
             "path":filename,
             "type":"plaintext","size":len(content)}]

def index_docx(filename):
    content = read_docx(filename)
    if not content or len(content) < 50:
        return []
    key = normalise_key(filename)
    meta = COOKBOOK_META.get(key)
    title, extra_kw = meta if meta else (filename.replace('.docx',''), '')
    kw = extract_keywords(content) + (' '+extra_kw if extra_kw else '')
    clean = re.sub(r'\s+',' ',content).strip()
    fn_l = filename.lower()
    cat = ('sales' if 'ready to sell' in fn_l else 'general')
    return [{"id":key,"title":title,"category":cat,"product":"CX Solutions",
             "kw":kw,"summary":clean[:300],
             "path":filename,
             "type":"plaintext","size":len(content)}]

def index_json_articles(filename):
    with open(filename,'r',encoding='utf-8') as f: data=json.load(f)
    product = data.get('meta',{}).get('product','QuestionPro')
    entries = []
    for art in data.get('articles',[]):
        art_id = art.get('id','')
        content = art.get('content','') or ''
        summary = art.get('summary','') or ''
        if art.get('word_count',0)<20 or len(content)<80: continue
        headings = ' '.join(h.get('text','') for h in (art.get('headings') or []))
        steps = ' '.join(s.get('text','') if isinstance(s,dict) else str(s) for s in (art.get('steps') or []))
        kw = extract_keywords(f"{art.get('title','')} {summary} {headings} {steps} {content}")
        eid = re.sub(r'[^a-z0-9\-_]','-',f"{normalise_key(filename)}-{art_id}".lower())
        entries.append({"id":eid,"title":art.get('title','Untitled'),
                        "category":'api' if 'API' in filename else 'help',
                        "product":product,"kw":kw,"summary":(summary or content)[:200],
                        "path":filename,
                        "article_id":art_id,"url":art.get('url',''),
                        "type":"json_article","word_count":art.get('word_count',0)})
    return entries

def main():
    index = []
    all_files = sorted(f for f in os.listdir('.') if f not in SKIP_FILES)
    txt_files  = [f for f in all_files if f.endswith('.txt')]
    docx_files = [f for f in all_files if f.endswith('.docx')]

    print(f"Indexing {len(txt_files)} .txt + {len(docx_files)} .docx files from repo root\n")

    for fn in txt_files:
        try:
            with open(fn,'r',encoding='utf-8',errors='replace') as f: peek=f.read(10).strip()
            if peek.startswith('{') or peek.startswith('['):
                entries = index_json_articles(fn)
                print(f"  {fn}: {len(entries)} articles (JSON)")
            else:
                entries = index_plaintext(fn)
                print(f"  {fn}: plaintext -> '{entries[0]['title']}'")
            index.extend(entries)
        except Exception as e:
            print(f"  {fn}: ERROR -- {e}")

    for fn in docx_files:
        try:
            entries = index_docx(fn)
            if entries:
                print(f"  {fn}: docx -> '{entries[0]['title']}'")
            else:
                print(f"  {fn}: docx -> skipped (no readable content)")
            index.extend(entries)
        except Exception as e:
            print(f"  {fn}: ERROR -- {e}")

    with open(OUTPUT_FILE,'w',encoding='utf-8') as f:
        json.dump(index,f,separators=(',',':'))
    print(f"\n✅ {OUTPUT_FILE}: {len(index)} entries, {os.path.getsize(OUTPUT_FILE)//1024} KB")

if __name__ == '__main__':
    main()

